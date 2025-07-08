import streamlit as st
import os
import tempfile
from typing import List, Dict, Any, Optional
import hashlib
from datetime import datetime
from pathlib import Path
import requests
import json

# Document processing - Windows compatible libraries
import PyPDF2
from docx import Document
import markdown
from bs4 import BeautifulSoup

# Text processing
from sentence_transformers import SentenceTransformer
import numpy as np

# Qdrant
from qdrant_client import QdrantClient
from qdrant_client.http import models as qdrant_models
from qdrant_client.http.models import Distance, VectorParams, PointStruct

# Environment and config
from dotenv import load_dotenv

# UI components
import plotly.graph_objects as go
import pandas as pd

# Load environment variables
load_dotenv()

# Configure page
st.set_page_config(
    page_title="AI Document Assistant",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for beautiful UI
st.markdown("""
<style>
    /* Main container styling */
    .main {
        padding: 2rem;
        background: linear-gradient(135deg, #0f0f1e 0%, #1a1a2e 100%);
    }
    
    /* Header styling */
    .main-header {
        text-align: center;
        margin-bottom: 3rem;
        animation: fadeIn 1s ease-in;
    }
    
    @keyframes fadeIn {
        from { opacity: 0; transform: translateY(-20px); }
        to { opacity: 1; transform: translateY(0); }
    }
    
    h1 {
        background: linear-gradient(120deg, #7c3aed, #ec4899, #f59e0b);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-size: 3.5rem;
        font-weight: 800;
        margin-bottom: 0.5rem;
    }
    
    .subtitle {
        color: #a1a1aa;
        font-size: 1.2rem;
        margin-bottom: 2rem;
    }
    
    /* Card styling */
    .stTabs [data-baseweb="tab-list"] {
        gap: 1rem;
        background-color: rgba(255, 255, 255, 0.05);
        padding: 1rem;
        border-radius: 15px;
        backdrop-filter: blur(10px);
    }
    
    .stTabs [data-baseweb="tab"] {
        background-color: transparent;
        border-radius: 10px;
        padding: 0.75rem 1.5rem;
        color: #a1a1aa;
        font-weight: 600;
        transition: all 0.3s ease;
    }
    
    .stTabs [data-baseweb="tab"]:hover {
        background-color: rgba(124, 58, 237, 0.2);
        color: #7c3aed;
    }
    
    .stTabs [aria-selected="true"] {
        background: linear-gradient(135deg, #7c3aed, #ec4899);
        color: white;
    }
    
    /* Upload area styling */
    .upload-area {
        background: rgba(255, 255, 255, 0.05);
        border: 2px dashed #7c3aed;
        border-radius: 20px;
        padding: 3rem;
        text-align: center;
        transition: all 0.3s ease;
        backdrop-filter: blur(10px);
    }
    
    .upload-area:hover {
        border-color: #ec4899;
        background: rgba(124, 58, 237, 0.1);
        transform: translateY(-2px);
        box-shadow: 0 10px 30px rgba(124, 58, 237, 0.3);
    }
    
    /* Chat messages */
    .chat-message {
        padding: 1.5rem;
        margin: 1rem 0;
        border-radius: 15px;
        animation: slideIn 0.3s ease;
        backdrop-filter: blur(10px);
    }
    
    @keyframes slideIn {
        from { opacity: 0; transform: translateX(-20px); }
        to { opacity: 1; transform: translateX(0); }
    }
    
    .user-message {
        background: linear-gradient(135deg, rgba(124, 58, 237, 0.2), rgba(236, 72, 153, 0.2));
        border-left: 4px solid #7c3aed;
        margin-left: 2rem;
    }
    
    .assistant-message {
        background: rgba(255, 255, 255, 0.05);
        border-left: 4px solid #ec4899;
        margin-right: 2rem;
    }
    
    /* Button styling */
    .stButton > button {
        background: linear-gradient(135deg, #7c3aed, #ec4899);
        color: white;
        border: none;
        padding: 0.75rem 2rem;
        border-radius: 30px;
        font-weight: 600;
        transition: all 0.3s ease;
        box-shadow: 0 4px 15px rgba(124, 58, 237, 0.4);
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 20px rgba(124, 58, 237, 0.6);
    }
    
    /* Sidebar styling */
    .css-1d391kg {
        background: #0f0f1e;
        border-right: 1px solid rgba(255, 255, 255, 0.1);
    }
    
    /* Select box styling */
    .stSelectbox > div > div {
        background: rgba(255, 255, 255, 0.05);
        border: 1px solid rgba(124, 58, 237, 0.5);
        border-radius: 10px;
    }
    
    /* Success/Error messages */
    .stSuccess {
        background: linear-gradient(135deg, rgba(16, 185, 129, 0.2), rgba(5, 150, 105, 0.2));
        border-left: 4px solid #10b981;
        border-radius: 10px;
        padding: 1rem;
    }
    
    .stError {
        background: linear-gradient(135deg, rgba(239, 68, 68, 0.2), rgba(220, 38, 38, 0.2));
        border-left: 4px solid #ef4444;
        border-radius: 10px;
        padding: 1rem;
    }
    
    /* Metrics styling */
    .metric-card {
        background: rgba(255, 255, 255, 0.05);
        border-radius: 15px;
        padding: 1.5rem;
        text-align: center;
        backdrop-filter: blur(10px);
        border: 1px solid rgba(255, 255, 255, 0.1);
        transition: all 0.3s ease;
    }
    
    .metric-card:hover {
        transform: translateY(-5px);
        box-shadow: 0 10px 30px rgba(124, 58, 237, 0.3);
        border-color: #7c3aed;
    }
    
    .metric-value {
        font-size: 2rem;
        font-weight: 700;
        background: linear-gradient(120deg, #7c3aed, #ec4899);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    
    .metric-label {
        color: #a1a1aa;
        font-size: 0.9rem;
        margin-top: 0.5rem;
    }
    
    /* Ollama status indicator */
    .ollama-status {
        padding: 0.5rem 1rem;
        border-radius: 20px;
        font-size: 0.9rem;
        font-weight: 600;
        display: inline-flex;
        align-items: center;
        gap: 0.5rem;
    }
    
    .ollama-online {
        background: rgba(16, 185, 129, 0.2);
        color: #10b981;
        border: 1px solid #10b981;
    }
    
    .ollama-offline {
        background: rgba(239, 68, 68, 0.2);
        color: #ef4444;
        border: 1px solid #ef4444;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'vector_store' not in st.session_state:
    st.session_state.vector_store = None
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'embedder' not in st.session_state:
    st.session_state.embedder = None
if 'documents_processed' not in st.session_state:
    st.session_state.documents_processed = 0
if 'total_chunks' not in st.session_state:
    st.session_state.total_chunks = 0
if 'qdrant_client' not in st.session_state:
    st.session_state.qdrant_client = None
if 'ollama_models' not in st.session_state:
    st.session_state.ollama_models = []

# Configuration class
class Config:
    EMBEDDING_MODEL = "sentence-transformers/all-MiniLM-L6-v2"
    CHUNK_SIZE = 1000
    CHUNK_OVERLAP = 200
    QDRANT_HOST = os.getenv("QDRANT_HOST", "localhost")
    QDRANT_PORT = int(os.getenv("QDRANT_PORT", 6333))
    COLLECTION_NAME = "document_qa"
    OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
    
    # Default Ollama models (will be updated from API)
    DEFAULT_OLLAMA_MODELS = [
        "llama3:latest",
        "llama2:latest",
        "mistral:latest",
        "codellama:latest",
        "phi:latest"
    ]

# Ollama integration
class OllamaClient:
    """Client for interacting with Ollama API"""
    
    def __init__(self, base_url: str = Config.OLLAMA_BASE_URL):
        self.base_url = base_url
        
    def is_available(self) -> bool:
        """Check if Ollama is running"""
        try:
            response = requests.get(f"{self.base_url}/api/tags")
            return response.status_code == 200
        except:
            return False
    
    def list_models(self) -> List[str]:
        """Get list of available models"""
        try:
            response = requests.get(f"{self.base_url}/api/tags")
            if response.status_code == 200:
                models = response.json().get('models', [])
                return [model['name'] for model in models]
        except:
            pass
        return Config.DEFAULT_OLLAMA_MODELS
    
    def generate(self, model: str, prompt: str, context: str = "") -> str:
        """Generate response from Ollama"""
        try:
            payload = {
                "model": model,
                "prompt": f"Context:\n{context}\n\nQuestion: {prompt}\n\nAnswer:",
                "stream": False,
                "options": {
                    "temperature": 0.7,
                    "num_predict": 2048
                }
            }
            
            response = requests.post(
                f"{self.base_url}/api/generate",
                json=payload,
                timeout=60
            )
            
            if response.status_code == 200:
                return response.json().get('response', 'Error: No response from model')
            else:
                return f"Error: {response.status_code} - {response.text}"
                
        except requests.exceptions.Timeout:
            return "Error: Request timed out. The model might be loading or the response is taking too long."
        except Exception as e:
            return f"Error: {str(e)}"
    
    def pull_model(self, model_name: str) -> bool:
        """Pull a model from Ollama registry"""
        try:
            response = requests.post(
                f"{self.base_url}/api/pull",
                json={"name": model_name},
                stream=True
            )
            return response.status_code == 200
        except:
            return False

# Document processing functions
class DocumentProcessor:
    """Handles document loading without langchain"""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> List[Dict[str, Any]]:
        """Extract text from PDF file"""
        documents = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text.strip():
                    documents.append({
                        'content': text,
                        'metadata': {
                            'page': page_num + 1,
                            'total_pages': len(pdf_reader.pages)
                        }
                    })
        
        return documents
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> List[Dict[str, Any]]:
        """Extract text from DOCX file"""
        documents = []
        doc = Document(file_path)
        
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        if full_text:
            documents.append({
                'content': '\n'.join(full_text),
                'metadata': {
                    'paragraphs': len(doc.paragraphs)
                }
            })
        
        return documents
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> List[Dict[str, Any]]:
        """Extract text from TXT file"""
        documents = []
        
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
            if content.strip():
                documents.append({
                    'content': content,
                    'metadata': {
                        'file_size': len(content)
                    }
                })
        
        return documents
    
    @staticmethod
    def extract_text_from_markdown(file_path: str) -> List[Dict[str, Any]]:
        """Extract text from Markdown file"""
        documents = []
        
        with open(file_path, 'r', encoding='utf-8') as file:
            md_content = file.read()
            
        # Convert markdown to HTML then extract text
        html = markdown.markdown(md_content)
        soup = BeautifulSoup(html, 'html.parser')
        text = soup.get_text()
        
        if text.strip():
            documents.append({
                'content': text,
                'metadata': {
                    'original_format': 'markdown'
                }
            })
        
        return documents

class TextSplitter:
    """Custom text splitter for chunking documents"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def split_text(self, text: str) -> List[str]:
        """Split text into chunks"""
        chunks = []
        
        # Split by paragraphs first
        paragraphs = text.split('\n\n')
        
        current_chunk = ""
        for para in paragraphs:
            # If adding this paragraph exceeds chunk size, save current chunk
            if len(current_chunk) + len(para) > self.chunk_size and current_chunk:
                chunks.append(current_chunk.strip())
                # Start new chunk with overlap
                overlap_text = current_chunk[-self.chunk_overlap:] if len(current_chunk) > self.chunk_overlap else current_chunk
                current_chunk = overlap_text + "\n\n" + para
            else:
                current_chunk += "\n\n" + para if current_chunk else para
        
        # Add the last chunk
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        # If chunks are still too large, split them further
        final_chunks = []
        for chunk in chunks:
            if len(chunk) > self.chunk_size:
                # Split by sentences
                sentences = chunk.split('. ')
                sub_chunk = ""
                for sentence in sentences:
                    if len(sub_chunk) + len(sentence) > self.chunk_size and sub_chunk:
                        final_chunks.append(sub_chunk.strip())
                        sub_chunk = sentence
                    else:
                        sub_chunk += ". " + sentence if sub_chunk else sentence
                if sub_chunk:
                    final_chunks.append(sub_chunk.strip())
            else:
                final_chunks.append(chunk)
        
        return final_chunks

# Initialize embedder
@st.cache_resource
def get_embedder():
    return SentenceTransformer(Config.EMBEDDING_MODEL)

# Initialize Qdrant client
@st.cache_resource
def get_qdrant_client():
    try:
        return QdrantClient(host=Config.QDRANT_HOST, port=Config.QDRANT_PORT)
    except Exception as e:
        st.error(f"Failed to connect to Qdrant: {e}")
        return None

# Initialize Ollama client
@st.cache_resource
def get_ollama_client():
    return OllamaClient(Config.OLLAMA_BASE_URL)

# Process documents
def process_documents(files) -> List[Dict[str, Any]]:
    """Process uploaded documents and split into chunks"""
    all_documents = []
    processor = DocumentProcessor()
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    for idx, file in enumerate(files):
        status_text.text(f"Processing {file.name}...")
        
        # Save uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file.name.split('.')[-1]}") as tmp_file:
            tmp_file.write(file.getbuffer())
            tmp_path = tmp_file.name
        
        try:
            # Extract text based on file type
            file_type = file.name.split('.')[-1].lower()
            
            if file_type == 'pdf':
                documents = processor.extract_text_from_pdf(tmp_path)
            elif file_type in ['docx', 'doc']:
                documents = processor.extract_text_from_docx(tmp_path)
            elif file_type == 'txt':
                documents = processor.extract_text_from_txt(tmp_path)
            elif file_type in ['md', 'markdown']:
                documents = processor.extract_text_from_markdown(tmp_path)
            else:
                st.warning(f"Unsupported file type: {file_type}")
                continue
            
            # Add source metadata
            for doc in documents:
                doc['metadata']['source'] = file.name
                doc['metadata']['upload_time'] = datetime.now().isoformat()
            
            all_documents.extend(documents)
            
        except Exception as e:
            st.error(f"Error processing {file.name}: {str(e)}")
        finally:
            os.unlink(tmp_path)
        
        progress_bar.progress((idx + 1) / len(files))
    
    status_text.text("Splitting documents into chunks...")
    
    # Split documents into chunks
    splitter = TextSplitter(chunk_size=Config.CHUNK_SIZE, chunk_overlap=Config.CHUNK_OVERLAP)
    chunks = []
    
    for doc in all_documents:
        text_chunks = splitter.split_text(doc['content'])
        for chunk_text in text_chunks:
            chunks.append({
                'content': chunk_text,
                'metadata': doc['metadata']
            })
    
    progress_bar.empty()
    status_text.empty()
    
    return chunks

def store_in_qdrant(chunks: List[Dict[str, Any]], embedder, client):
    """Store document chunks in Qdrant"""
    with st.spinner("Creating embeddings and storing in Qdrant..."):
        # Get embeddings
        texts = [chunk['content'] for chunk in chunks]
        embeddings = embedder.encode(texts, show_progress_bar=True)
        
        # Create collection if it doesn't exist
        try:
            client.get_collection(Config.COLLECTION_NAME)
        except:
            client.create_collection(
                collection_name=Config.COLLECTION_NAME,
                vectors_config=VectorParams(
                    size=384,  # all-MiniLM-L6-v2 dimension
                    distance=Distance.COSINE
                )
            )
        
        # Prepare points for Qdrant
        points = []
        for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
            points.append(
                PointStruct(
                    id=i,
                    vector=embedding.tolist(),
                    payload={
                        "text": chunk['content'],
                        "source": chunk['metadata'].get('source', 'Unknown'),
                        "metadata": chunk['metadata']
                    }
                )
            )
        
        # Upload to Qdrant in batches
        batch_size = 100
        for i in range(0, len(points), batch_size):
            batch = points[i:i + batch_size]
            client.upsert(
                collection_name=Config.COLLECTION_NAME,
                points=batch
            )
        
        return True

def search_qdrant(query: str, embedder, client, k: int = 5):
    """Search Qdrant for relevant documents"""
    # Get query embedding
    query_embedding = embedder.encode(query).tolist()
    
    # Search
    results = client.search(
        collection_name=Config.COLLECTION_NAME,
        query_vector=query_embedding,
        limit=k
    )
    
    return results

# Main UI
def main():
    # Header
    st.markdown("""
    <div class="main-header">
        <h1>AI Document Assistant</h1>
        <p class="subtitle">Upload documents, ask questions, get intelligent answers - Powered by Local LLMs</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Initialize components
    if not st.session_state.embedder:
        st.session_state.embedder = get_embedder()
    
    if not st.session_state.qdrant_client:
        st.session_state.qdrant_client = get_qdrant_client()
    
    # Initialize Ollama client
    ollama_client = get_ollama_client()
    
    # Check Ollama status
    ollama_available = ollama_client.is_available()
    
    # Sidebar
    with st.sidebar:
        st.markdown("### ⚙️ Configuration")
        
        # Ollama status
        if ollama_available:
            st.markdown('<div class="ollama-status ollama-online">🟢 Ollama Connected</div>', unsafe_allow_html=True)
            
            # Get available models
            available_models = ollama_client.list_models()
            st.session_state.ollama_models = available_models
            
            # Model selection
            selected_model = st.selectbox(
                "Select Model",
                available_models,
                help="Choose the local LLM model for answering questions"
            )
        else:
            st.markdown('<div class="ollama-status ollama-offline">🔴 Ollama Offline</div>', unsafe_allow_html=True)
            st.error("Ollama is not running!")
            st.code("# Start Ollama:\nollama serve", language="bash")
            
            # Show default models anyway
            selected_model = st.selectbox(
                "Select Model (when Ollama is running)",
                Config.DEFAULT_OLLAMA_MODELS,
                disabled=True
            )
        
        st.markdown("---")
        
        # Ollama settings
        with st.expander("🔧 Ollama Settings"):
            ollama_url = st.text_input(
                "Ollama Base URL",
                value=Config.OLLAMA_BASE_URL,
                help="Default: http://localhost:11434"
            )
            
            if st.button("Test Connection"):
                test_client = OllamaClient(ollama_url)
                if test_client.is_available():
                    st.success("✅ Connected to Ollama!")
                else:
                    st.error("❌ Cannot connect to Ollama")
            
            st.markdown("### Pull New Model")
            new_model = st.text_input("Model name (e.g., llama3:latest)")
            if st.button("Pull Model") and new_model:
                with st.spinner(f"Pulling {new_model}..."):
                    if ollama_client.pull_model(new_model):
                        st.success(f"✅ Successfully pulled {new_model}")
                        st.rerun()
                    else:
                        st.error("Failed to pull model")
        
        st.markdown("---")
        
        # Statistics
        if st.session_state.documents_processed > 0:
            st.markdown("### 📊 Statistics")
            
            col1, col2 = st.columns(2)
            with col1:
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-value">{st.session_state.documents_processed}</div>
                    <div class="metric-label">Documents</div>
                </div>
                """, unsafe_allow_html=True)
            
            with col2:
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-value">{st.session_state.total_chunks}</div>
                    <div class="metric-label">Chunks</div>
                </div>
                """, unsafe_allow_html=True)
            
            # Visualization
            if st.button("📈 Show Analytics"):
                fig = go.Figure(data=[
                    go.Bar(
                        x=['Documents', 'Chunks', 'Queries'],
                        y=[st.session_state.documents_processed, 
                           st.session_state.total_chunks,
                           len(st.session_state.chat_history) // 2],
                        marker_color=['#7c3aed', '#ec4899', '#f59e0b']
                    )
                ])
                fig.update_layout(
                    title="System Analytics",
                    showlegend=False,
                    plot_bgcolor='rgba(0,0,0,0)',
                    paper_bgcolor='rgba(0,0,0,0)',
                    font=dict(color='white')
                )
                st.plotly_chart(fig, use_container_width=True)
        
        st.markdown("---")
        
        # Instructions
        with st.expander("📖 How to use"):
            st.markdown("""
            1. **Install Ollama**: Download from [ollama.ai](https://ollama.ai)
            2. **Start Ollama**: Run `ollama serve` in terminal
            3. **Pull a model**: e.g., `ollama pull llama3`
            4. **Upload Documents**: Go to Upload tab
            5. **Process**: Click Process Documents
            6. **Ask Questions**: Use the Chat tab
            
            **Supported formats**: PDF, TXT, DOCX, MD
            """)
        
        # Clear data
        if st.button("🗑️ Clear All Data", type="secondary", use_container_width=True):
            st.session_state.vector_store = None
            st.session_state.chat_history = []
            st.session_state.documents_processed = 0
            st.session_state.total_chunks = 0
            st.rerun()
    
    # Main content
    tabs = st.tabs(["📁 Upload Documents", "💬 Chat", "📚 Document Library"])
    
    # Upload tab
    with tabs[0]:
        st.markdown("""
        <div class="upload-area">
            <h3>📤 Upload Your Documents</h3>
            <p>Support for PDF, TXT, DOCX, and Markdown files</p>
        </div>
        """, unsafe_allow_html=True)
        
        uploaded_files = st.file_uploader(
            "Choose files",
            type=['pdf', 'txt', 'docx', 'doc', 'md', 'markdown'],
            accept_multiple_files=True,
            label_visibility="collapsed"
        )
        
        if uploaded_files:
            col1, col2, col3 = st.columns([1, 2, 1])
            with col2:
                if st.button("🚀 Process Documents", use_container_width=True):
                    if not st.session_state.qdrant_client:
                        st.error("Cannot connect to Qdrant! Make sure it's running on localhost:6333")
                        st.code("docker run -p 6333:6333 qdrant/qdrant", language="bash")
                        return
                    
                    # Process documents
                    chunks = process_documents(uploaded_files)
                    
                    if chunks:
                        # Store in Qdrant
                        success = store_in_qdrant(
                            chunks,
                            st.session_state.embedder,
                            st.session_state.qdrant_client
                        )
                        
                        if success:
                            st.session_state.documents_processed += len(uploaded_files)
                            st.session_state.total_chunks += len(chunks)
                            st.session_state.vector_store = True
                            
                            st.success(f"""
                            ✅ Successfully processed {len(uploaded_files)} documents into {len(chunks)} chunks!
                            You can now ask questions in the Chat tab.
                            """)
                            
                            # Show processed files
                            st.markdown("### Processed Files:")
                            for file in uploaded_files:
                                st.markdown(f"- 📄 {file.name}")
    
    # Chat tab
    with tabs[1]:
        if not st.session_state.vector_store:
            st.info("📤 Please upload and process documents first in the Upload Documents tab.")
        elif not ollama_available:
            st.warning("🔴 Ollama is not running. Please start Ollama to use the chat feature.")
            st.code("ollama serve", language="bash")
        else:
            # Chat interface
            st.markdown("### 💬 Ask Questions About Your Documents")
            
            # Display chat history
            for i in range(0, len(st.session_state.chat_history), 2):
                if i < len(st.session_state.chat_history):
                    # User message
                    user_msg = st.session_state.chat_history[i]
                    st.markdown(f"""
                    <div class="chat-message user-message">
                        <strong>You:</strong> {user_msg['content']}
                    </div>
                    """, unsafe_allow_html=True)
                    
                    # Assistant message
                    if i + 1 < len(st.session_state.chat_history):
                        assistant_msg = st.session_state.chat_history[i + 1]
                        st.markdown(f"""
                        <div class="chat-message assistant-message">
                            <strong>Assistant:</strong> {assistant_msg['content']}
                        </div>
                        """, unsafe_allow_html=True)
                        
                        if assistant_msg.get('sources'):
                            with st.expander("📚 Sources"):
                                for source in assistant_msg['sources']:
                                    st.markdown(f"- {source}")
            
            # Query input
            query = st.text_input(
                "Ask a question:",
                placeholder="What would you like to know about your documents?",
                key="query_input"
            )
            
            col1, col2, col3 = st.columns([1, 1, 3])
            with col1:
                if st.button("🔍 Ask", use_container_width=True):
                    if query:
                        with st.spinner("Searching and thinking..."):
                            # Search Qdrant
                            results = search_qdrant(
                                query,
                                st.session_state.embedder,
                                st.session_state.qdrant_client,
                                k=5
                            )
                            
                            if results:
                                # Prepare context
                                context_parts = []
                                sources = set()
                                
                                for i, result in enumerate(results):
                                    context_parts.append(f"[{i+1}] {result.payload['text']}")
                                    sources.add(result.payload['source'])
                                
                                context = "\n\n".join(context_parts)
                                
                                # Get response from Ollama
                                with st.spinner(f"Generating response with {selected_model}..."):
                                    response = ollama_client.generate(
                                        model=selected_model,
                                        prompt=query,
                                        context=context
                                    )
                                
                                # Add to chat history
                                st.session_state.chat_history.append({
                                    'role': 'user',
                                    'content': query
                                })
                                
                                st.session_state.chat_history.append({
                                    'role': 'assistant',
                                    'content': response,
                                    'sources': list(sources)
                                })
                                
                                st.rerun()
                            else:
                                st.warning("No relevant documents found for your query.")
            
            with col2:
                if st.button("🔄 Clear Chat", use_container_width=True):
                    st.session_state.chat_history = []
                    st.rerun()
    
    # Document Library tab
    with tabs[2]:
        st.markdown("### 📚 Document Library")
        
        if st.session_state.vector_store:
            st.info(f"""
            📊 **Current Status:**
            - Total documents processed: {st.session_state.documents_processed}
            - Total chunks in database: {st.session_state.total_chunks}
            - Vector database: Qdrant
            - Embedding model: {Config.EMBEDDING_MODEL}
            - LLM: Local Ollama models
            """)
            
            # Add search functionality
            search_query = st.text_input("🔍 Search documents:", placeholder="Enter keywords...")
            if search_query and st.button("Search", key="search_docs"):
                with st.spinner("Searching..."):
                    results = search_qdrant(
                        search_query,
                        st.session_state.embedder,
                        st.session_state.qdrant_client,
                        k=10
                    )
                    
                    if results:
                        st.markdown("### Search Results:")
                        for i, result in enumerate(results, 1):
                            score = result.score
                            with st.expander(f"Result {i} - {result.payload['source']} (Score: {score:.2f})"):
                                st.write(result.payload['text'][:500] + "..." if len(result.payload['text']) > 500 else result.payload['text'])
                                st.markdown(f"**Source:** {result.payload['source']}")
                                if 'page' in result.payload['metadata']:
                                    st.markdown(f"**Page:** {result.payload['metadata']['page']}")
                    else:
                        st.info("No results found.")
        else:
            st.info("No documents uploaded yet. Please upload documents in the Upload Documents tab.")
    
    # Footer
    st.markdown("---")
    st.markdown(
        "<center>Built with ❤️ using Streamlit, Qdrant, and Ollama</center>",
        unsafe_allow_html=True
    )

if __name__ == "__main__":
    main()